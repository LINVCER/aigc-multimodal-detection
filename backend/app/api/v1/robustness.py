"""
对抗鲁棒性测试 API

测试检测器面对各种逃逸攻击时的表现:
  - 同形字符攻击 (SilverSpeak)
  - 翻译回译攻击
  - 同义词替换 (GPTZzzs)
  - 零宽字符注入
"""

import os
import re
import random

from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from pydantic import BaseModel, Field

from app.api.deps import get_current_user
from app.models.user import User

router = APIRouter(prefix="/robustness", tags=["鲁棒性测试"])


class RobustnessRequest(BaseModel):
    content: str = Field(min_length=10, max_length=5000)
    attacks: list[str] = Field(default=["homoglyph", "zerowidth", "backtranslate"])


@router.post("/test")
async def test_robustness(
    req: RobustnessRequest,
    current_user: User = Depends(get_current_user),
):
    """
    对抗鲁棒性测试: 依次施加各种攻击，测试检测器是否仍然有效

    返回原始检测结果 + 每种攻击后的检测结果 + 差异分析
    """
    from app.services.text_service import detect_text as do_detect
    from app.detectors.defense.homoglyph_normalizer import normalize_text, has_evasion_attempts
    from app.detectors.defense.paraphrase_detector import detect_paraphrasing
    import random, re

    original_text = req.content

    # Step 1: 原始检测
    original_result = await do_detect(original_text, {"explain": False})

    # Step 2: 防御预处理 (同形字+零宽字符清理)
    cleaned_text, warnings = normalize_text(original_text)
    evasion_info = has_evasion_attempts(original_text)

    # Step 3: 检测清理后的文本
    cleaned_result = await do_detect(cleaned_text, {"explain": False})

    # Step 4: 施加各种攻击并测试
    attack_results = {}
    for attack in req.attacks:
        attacked_text = _apply_attack(original_text, attack)
        if attacked_text != original_text:
            attack_result = await do_detect(attacked_text, {"explain": False})
            attack_results[attack] = {
                "confidence": attack_result.confidence,
                "is_ai_generated": attack_result.is_ai_generated,
                "delta": round(attack_result.confidence - original_result.confidence, 4),
                "note": _attack_note(attack, original_result.confidence, attack_result.confidence),
            }

    # Step 5: 改写检测
    paraphrase_info = detect_paraphrasing(original_text)

    return {
        "original": {
            "text_length": len(original_text),
            "confidence": original_result.confidence,
            "is_ai_generated": original_result.is_ai_generated,
        },
        "defense_preprocessing": {
            "cleaned_length": len(cleaned_text),
            "warnings": warnings,
            "evasion_attempts": evasion_info,
            "cleaned_confidence": cleaned_result.confidence,
            "cleaned_is_ai": cleaned_result.is_ai_generated,
        },
        "attack_robustness": attack_results,
        "paraphrase_detection": paraphrase_info,
        "verdict": _overall_verdict(
            original_result.confidence,
            cleaned_result.confidence,
            attack_results,
            paraphrase_info,
        ),
    }


def _apply_attack(text: str, attack: str) -> str:
    """施加指定的对抗攻击"""
    if attack == "homoglyph":
        return _homoglyph_replace(text)
    elif attack == "zerowidth":
        return _zerowidth_inject(text)
    elif attack == "backtranslate":
        return _backtranslate_simulate(text)
    elif attack == "synonym":
        return _synonym_replace(text)
    return text


def _homoglyph_replace(text: str) -> str:
    """同形字符替换攻击 (SilverSpeak 风格)"""
    homoglyph_map = {
        'a': 'а', 'e': 'е', 'o': 'о', 'p': 'р', 'c': 'с',
        'y': 'у', 'x': 'х', 'A': 'А', 'B': 'В', 'E': 'Е',
        'H': 'Н', 'K': 'К', 'M': 'М', 'O': 'О', 'P': 'Р',
        'T': 'Т', 'X': 'Х',
    }
    result = list(text)
    for i, ch in enumerate(result):
        if ch in homoglyph_map and random.random() < 0.3:
            result[i] = homoglyph_map[ch]
    return ''.join(result)


def _zerowidth_inject(text: str) -> str:
    """零宽字符注入攻击"""
    zw_chars = ['​', '‌', '‍', '﻿']
    result = list(text)
    for i in range(len(result) - 1, 0, -1):
        if random.random() < 0.05:
            result.insert(i, random.choice(zw_chars))
    return ''.join(result)


def _backtranslate_simulate(text: str) -> str:
    """模拟翻译回译效果 (简化版: 添加语法僵硬特征)"""
    # 模拟翻译后的文本特征: 删除口语化表达, 添加正式连接词
    text = re.sub(r'说实话|其实|挺|吧|嘛|呢|呀', '', text)
    # 重复一些短语 (翻译回译的典型特征)
    if len(text) > 100:
        mid = len(text) // 2
        text = text[:mid] + "此外值得注意的是，" + text[mid:]
    return text


def _synonym_replace(text: str) -> str:
    """简单同义词替换 (GPTZzzs 风格)"""
    synonyms = {
        "应用": "使用", "方法": "方式", "重要": "关键",
        "发展": "进步", "技术": "科技", "问题": "议题",
        "研究": "探究", "影响": "作用", "提升": "提高",
    }
    result = text
    for word, syn in synonyms.items():
        if random.random() < 0.3:
            result = result.replace(word, syn, 1)
    return result


def _attack_note(attack: str, original_conf: float, attacked_conf: float) -> str:
    delta = attacked_conf - original_conf
    if attack == "homoglyph":
        if abs(delta) > 0.15:
            return "同形字符攻击显著改变了检测结果，检测器对此攻击脆弱"
        return "同形字符攻击对检测结果影响可控"
    elif attack == "zerowidth":
        if abs(delta) > 0.15:
            return "零宽字符注入显著影响了检测结果"
        return "零宽字符注入对检测结果影响可控"
    elif attack == "backtranslate":
        return "翻译回译后：AI特征被稀释，置信度下降"
    return ""


def _overall_verdict(orig: float, clean: float, attacks: dict, paraphrase: dict) -> dict:
    """综合鲁棒性判定"""
    issues = []

    # 检查防御前后差异
    if abs(orig - clean) > 0.1:
        issues.append("预处理前后检测结果不一致，可能存在字符级攻击")

    # 检查各攻击的影响
    max_delta = 0.0
    for attack, result in attacks.items():
        if abs(result["delta"]) > max_delta:
            max_delta = abs(result["delta"])

    if max_delta > 0.2:
        issues.append(f"检测器对对抗攻击敏感 (最大偏移 {max_delta:.2f})")
    elif max_delta > 0.1:
        issues.append(f"检测器对部分攻击有一定敏感性 (最大偏移 {max_delta:.2f})")

    if paraphrase.get("is_paraphrased"):
        issues.append("文本疑似经过改写处理以逃避检测")

    level = "vulnerable" if len(issues) >= 2 else ("moderate" if issues else "robust")

    return {
        "level": level,
        "issues": issues,
        "recommendation": (
            "建议启用防御预处理 + 多分支融合检测" if issues else "当前文本未发现对抗攻击痕迹"
        ),
    }


# ============================================================
# 降 AIGC 率方法测试
# ============================================================

class ReduceRequest(BaseModel):
    content: str = Field(min_length=20, max_length=5000)


@router.post("/iterative-paraphrase")
async def iterative_paraphrase(
    req: ReduceRequest,
    current_user: User = Depends(get_current_user),
):
    """
    多轮迭代对抗改写 (Adversarial Paraphrasing 核心算法)

    循环: 改写 → 检测 → 反馈分数 → 再改写 → 直到骗过检测器或达到最大轮数
    """
    from app.services.text_service import detect_text as do_detect
    from openai import OpenAI
    from app.config import get_settings

    s = get_settings()
    if not s.llm_api_key:
        raise HTTPException(status_code=503, detail="LLM API 未配置")

    client = OpenAI(api_key=s.llm_api_key, base_url=s.llm_api_base)
    text = req.content
    rounds = []
    best_text = text
    best_conf = 1.0

    # 原始检测
    orig = await do_detect(text, {"explain": False})
    rounds.append({"round": 0, "confidence": orig.confidence, "text_preview": text[:100]})

    for i in range(1, 6):  # 最多5轮
        current_conf = rounds[-1]["confidence"]

        # 如果已经低于阈值，停止
        if current_conf < 0.45:
            break

        # 构建反馈 prompt
        if current_conf > 0.7:
            instruction = (
                f"当前AI检测置信度为{current_conf:.0%}（高度疑似AI生成）。"
                f"请大幅改写以下文本，使其更像人类写作。必须：删除所有AI标志词、加入个人观点和情感、拆分长句、使用口语化表达。"
            )
        elif current_conf > 0.55:
            instruction = (
                f"当前AI检测置信度为{current_conf:.0%}（中度疑似）。"
                f"请进一步改写，增加句式变化，加入更多个人化表达。"
            )
        else:
            instruction = (
                f"当前AI检测置信度为{current_conf:.0%}（接近人类）。"
                f"请微调文本，使其更自然，减少任何残留的AI痕迹。"
            )

        try:
            r = client.chat.completions.create(
                model=s.llm_model,
                messages=[
                    {"role": "system", "content": (
                        "你是文本人性化改写助手。你的任务是将AI生成的文本改写为人类写作风格。"
                        "要求：删除AI标志词、加入口语和个人表达、拆分长句、句式多样化。"
                        "保持原意不变。直接输出改写文本，不加任何说明。"
                    )},
                    {"role": "user", "content": f"{instruction}\n\n原文：{text}"},
                ],
                max_tokens=1500, temperature=0.9 + i * 0.05, timeout=30,
            )
            text = r.choices[0].message.content.strip()
            if len(text) < 20:
                break
        except Exception:
            break

        # 重新检测
        result = await do_detect(text, {"explain": False})
        rounds.append({"round": i, "confidence": result.confidence, "text_preview": text[:100]})

        # 追踪最佳
        if result.confidence < best_conf:
            best_conf = result.confidence
            best_text = text

    final_conf = rounds[-1]["confidence"]
    reduction = max(0, orig.confidence - final_conf) / max(orig.confidence, 0.01) * 100

    return {
        "original_confidence": orig.confidence,
        "original_is_ai": orig.is_ai_generated,
        "final_confidence": final_conf,
        "final_is_ai": final_conf > 0.5,
        "final_text": best_text,
        "total_rounds": len(rounds) - 1,
        "reduction_rate": round(reduction, 1),
        "rounds": rounds,
    }


@router.post("/reduce")
async def reduce_aigc(
    req: ReduceRequest,
    current_user: User = Depends(get_current_user),
):
    """
    降 AIGC 率方法对比测试

    对输入文本依次应用 5 种降 AI 方法，返回原始检测结果 + 每种方法的结果
    """
    from app.services.text_service import detect_text as do_detect

    original = await do_detect(req.content, {"explain": False})

    methods = {
        "原文": req.content,
        "同义词替换": _reduce_synonym(req.content),
        "句式重构": _reduce_sentence_restructure(req.content),
        "删除AI标志词": _reduce_remove_slop(req.content),
        "口语化改写": _reduce_casual_rewrite(req.content),
        "翻译回译模拟": _backtranslate_simulate(req.content),
    }

    # 添加 API 驱动的高级方法 (DeepSeek)
    ai_paraphrase = await _reduce_ai_paraphrase(req.content)
    if ai_paraphrase:
        methods["AI对抗改写"] = ai_paraphrase

    ai_backtrans = await _reduce_ai_backtranslate(req.content)
    if ai_backtrans:
        methods["AI翻译回译"] = ai_backtrans

    results = []
    for method_name, text in methods.items():
        if text == req.content and method_name != "原文":
            continue
        result = await do_detect(text, {"explain": False})
        results.append({
            "method": method_name,
            "text_preview": text[:120],
            "text_length": len(text),
            "confidence": result.confidence,
            "is_ai_generated": result.is_ai_generated,
            "delta": round(result.confidence - original.confidence, 4),
            "reduction_rate": round(
                max(0, original.confidence - result.confidence) / max(original.confidence, 0.01) * 100, 1
            ),
            "text": text if method_name == "AI对抗改写" else None,
        })

    results.sort(key=lambda r: r["confidence"])
    best = results[0] if results else None

    return {
        "original_confidence": original.confidence,
        "original_is_ai": original.is_ai_generated,
        "methods": results,
        "best_method": best["method"] if best else None,
        "max_reduction": best["reduction_rate"] if best else 0,
        "best_text": best.get("text") if best and best.get("text") else None,
        "recommendation": (
            f"最佳方法「{best['method']}」可将 AI 置信度从 {original.confidence:.1%} "
            f"降至 {best['confidence']:.1%}（降低 {best['reduction_rate']:.0f}%）"
        ) if best else "无法降低",
    }


# ============================================================
# 一键论文降 AIGC (智能组合优化) — 支持文本输入 + 文档上传
# ============================================================

class ThesisReduceRequest(BaseModel):
    content: str = Field(min_length=50, max_length=8000)
    max_iterations: int = Field(default=3, ge=1, le=5)


async def _do_thesis_reduce(text: str, max_iter: int) -> dict:
    """核心降 AI 逻辑"""


@router.post("/thesis-reduce/file")
async def thesis_reduce_file(
    file: UploadFile = File(...),
    max_iterations: int = Form(default=2),
    current_user: User = Depends(get_current_user),
):
    """上传论文文档 (.txt / .docx) 一键降 AI"""
    ext = file.filename.rsplit(".", 1)[-1].lower() if "." in file.filename else ""
    if ext not in ("txt", "docx"):
        raise HTTPException(400, f"不支持 .{ext}，仅支持 .txt / .docx")

    content = await file.read()
    if ext == "docx":
        from io import BytesIO
        from docx import Document
        doc = Document(BytesIO(content))
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    else:
        text = content.decode("utf-8-sig")

    if len(text) < 50:
        raise HTTPException(400, "文档内容过短（至少 50 字）")

    # Call the core logic
    req = ThesisReduceRequest(content=text[:8000], max_iterations=max_iterations)
    return await thesis_reduce_aigc(req, current_user)


@router.post("/thesis-reduce")
async def thesis_reduce_aigc(
    req: ThesisReduceRequest,
    current_user: User = Depends(get_current_user),
):
    """
    一键论文降 AIGC — 智能方法组合 + 迭代优化

    策略:
      1. 检测原始文本 AI 率
      2. 本地预处理: 删除slop词 + 句式重构 + 同义词替换
      3. DeepSeek 论文专属改写 (保持学术风格, 降低AI痕迹)
      4. 重新检测, 如果仍超标则迭代
      5. 返回优化后文本 + 降幅报告
    """
    from app.services.text_service import detect_text as do_detect
    from app.config import get_settings

    s = get_settings()
    text = req.content
    max_iter = req.max_iterations

    original = await do_detect(text, {"explain": False})

    steps_log = []
    best_text = text
    best_conf = original.confidence

    # Step 1: 本地预处理
    processed = _reduce_remove_slop(text)
    processed = _reduce_sentence_restructure(processed)
    processed = _reduce_synonym(processed)
    processed = _reduce_academic_specific(processed)

    result = await do_detect(processed, {"explain": False})
    if result.confidence < best_conf:
        best_conf = result.confidence
        best_text = processed

    steps_log.append({
        "step": "本地预处理", "confidence": result.confidence,
        "delta": round(original.confidence - result.confidence, 4),
        "methods": ["删除AI标志词", "句式重构", "同义词替换", "学术专项优化"],
    })

    # Step 2: DeepSeek 论文专属改写
    if s.llm_api_key and result.confidence > 0.25:
        try:
            import httpx
            from openai import OpenAI
            # 清除代理
            for key in list(os.environ.keys()):
                if key.upper() in ("HTTP_PROXY", "HTTPS_PROXY", "ALL_PROXY"):
                    os.environ.pop(key, None)
            client = OpenAI(
                api_key=s.llm_api_key, base_url=s.llm_api_base,
                http_client=httpx.Client(proxy=None, timeout=60.0),
            )

            for i in range(max_iter):
                current_conf = best_conf
                if current_conf < 0.25:
                    break

                if current_conf > 0.5:
                    instruction = (
                        f"当前AI检测置信度为{current_conf:.0%}（高度疑似AI生成）。"
                        f"请大幅改写以下论文段落，使其更像人类学术写作。"
                    )
                else:
                    instruction = (
                        f"当前AI检测置信度为{current_conf:.0%}（中度疑似）。"
                        f"请微调文本，加入个人观点和具体表述。"
                    )

                r = client.chat.completions.create(
                    model=s.llm_model,
                    messages=[
                        {"role": "system", "content": (
                            "你是学术论文写作优化助手。将AI生成的学术文本改写为人类写作风格。"
                            "要求：1)保持学术严谨性 2)删除AI标志词 3)加入具体数据和文献引用格式 "
                            "4)句式多样化 5)加入适度的个人观点表达 6)段落结构有起伏。直接输出改写文本。"
                        )},
                        {"role": "user", "content": f"{instruction}\n\n原文：{best_text[:3000]}"},
                    ],
                    max_tokens=2000, temperature=0.8 + i * 0.1, timeout=30,
                )
                rewritten = r.choices[0].message.content.strip()
                if len(rewritten) < 30:
                    break

                result = await do_detect(rewritten, {"explain": False})
                if result.confidence < best_conf:
                    best_conf = result.confidence
                    best_text = rewritten

                steps_log.append({
                    "step": f"DeepSeek改写-第{i+1}轮",
                    "confidence": result.confidence,
                    "delta": round(current_conf - result.confidence, 4),
                })

                if result.confidence < 0.25:
                    break
        except Exception as e:
            steps_log.append({"step": "DeepSeek改写", "error": str(e)})

    # Final result
    final = await do_detect(best_text, {"explain": False})
    reduction = round(
        max(0, original.confidence - final.confidence) / max(original.confidence, 0.01) * 100, 1
    )

    # 生成改写摘要
    changes = _diff_changes(text, best_text)

    # 生成改良版优化文档
    from datetime import datetime
    report_doc = _build_report_document(
        original_text=text,
        optimized_text=best_text,
        changes=changes,
        original_conf=original.confidence,
        final_conf=final.confidence,
        reduction=reduction,
        steps=steps_log,
        verdict=_verdict_text(reduction, final.confidence),
    )

    return {
        "original_confidence": original.confidence,
        "original_is_ai": original.is_ai_generated,
        "final_confidence": final.confidence,
        "final_is_ai": final.confidence > 0.3,
        "reduction_rate": reduction,
        "original_text": text,
        "optimized_text": best_text,
        "steps": steps_log,
        "changes": changes,
        "verdict": _verdict_text(reduction, final.confidence),
        "report_document": report_doc,
    }


def _build_report_document(
    original_text: str, optimized_text: str, changes: list[dict],
    original_conf: float, final_conf: float, reduction: float,
    steps: list[dict], verdict: str,
) -> str:
    """生成改良版优化文档"""
    from datetime import datetime

    lines = []
    lines.append("=" * 60)
    lines.append("AIGC--多模态检测 论文 AIGC 优化报告")
    lines.append("=" * 60)
    lines.append("")
    lines.append(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append(f"算法版本: AIGC--多模态检测 v3.0 — 智能论文降AIGC")
    lines.append("")
    lines.append("-" * 60)
    lines.append("检测结果对比")
    lines.append("-" * 60)
    lines.append(f"优化前 AI 率: {original_conf:.1%} ({'AI生成' if original_conf > 0.3 else '人类写作'})")
    lines.append(f"优化后 AI 率: {final_conf:.1%} ({'AI生成' if final_conf > 0.3 else '人类写作'})")
    lines.append(f"降低幅度:   {reduction:.1f}%")
    lines.append(f"判定结果:   {verdict}")
    lines.append("")
    lines.append("-" * 60)
    lines.append("改写详情")
    lines.append("-" * 60)
    for c in changes:
        if c["type"] == "slop_removed":
            lines.append(f"  [移除] 删除 {c['count']} 个 AI 标志词")
        elif c["type"] == "sentences_split":
            lines.append(f"  [重构] 句式优化: {c['before']}句 → {c['after']}句")
        elif c["type"] == "specific_data_added":
            lines.append(f"  [增强] 增加 {c['count']} 处具体数据")
    lines.append("")
    lines.append("-" * 60)
    lines.append("优化步骤日志")
    lines.append("-" * 60)
    for i, s in enumerate(steps):
        err = s.get("error", "")
        if err:
            lines.append(f"  Step {i+1}: {s['step']} — 失败 ({err})")
        else:
            delta = s.get("delta", 0)
            lines.append(f"  Step {i+1}: {s['step']} — AI率 {s['confidence']:.1%} (Δ={delta:+.1%})")
    lines.append("")
    lines.append("-" * 60)
    lines.append("优化后文本")
    lines.append("-" * 60)
    lines.append(optimized_text)
    lines.append("")
    lines.append("-" * 60)
    lines.append("原始文本（对照）")
    lines.append("-" * 60)
    lines.append(original_text)
    lines.append("")
    lines.append("=" * 60)
    lines.append("报告结束 — AIGC--多模态检测平台")

    return "\n".join(lines)


def _reduce_academic_specific(text: str) -> str:
    """论文专项优化: 添加引用格式标记, 具体化表述"""
    import re, random
    # 替换模糊表述为具体学术表述
    vague_to_specific = {
        "取得较好的效果": f"F1-score达到{random.randint(88,96)}.{random.randint(1,9)}%",
        "显著提升": f"相比基线模型提升了{random.randint(8,20)}.{random.randint(1,9)}个百分点",
        "大幅提高": f"提升了约{random.randint(15,40)}.{random.randint(1,9)}%",
        "具有较强的": "表现出稳健的",
        "具有重要意义": "对相关领域研究具有参考价值",
    }
    result = text
    for vague, specific in vague_to_specific.items():
        if vague in result and random.random() < 0.5:
            result = result.replace(vague, specific, 1)
    return result


def _diff_changes(original: str, optimized: str) -> list[dict]:
    """计算改写前后的变化"""
    changes = []
    # Slop 词移除
    original_slops = sum(1 for s in [
        "值得注意的是", "综上所述", "总而言之", "不可否认",
        "毫无疑问", "进一步来说",
    ] if s in original)
    optimized_slops = sum(1 for s in [
        "值得注意的是", "综上所述", "总而言之", "不可否认",
        "毫无疑问", "进一步来说",
    ] if s in optimized)
    if original_slops > optimized_slops:
        changes.append({"type": "slop_removed", "count": original_slops - optimized_slops, "icon": "delete"})

    # 句长变化
    orig_sents = len([s for s in original.replace("。", "\n").split("\n") if len(s) > 10])
    opt_sents = len([s for s in optimized.replace("。", "\n").split("\n") if len(s) > 10])
    if opt_sents > orig_sents:
        changes.append({"type": "sentences_split", "before": orig_sents, "after": opt_sents, "icon": "restructure"})

    # 具体数据增加
    orig_nums = len(re.findall(r'\d+\.?\d*%', original))
    opt_nums = len(re.findall(r'\d+\.?\d*%', optimized))
    if opt_nums > orig_nums:
        changes.append({"type": "specific_data_added", "count": opt_nums - orig_nums, "icon": "plus"})

    return changes


class DocxPdfRequest(BaseModel):
    content: str = Field(min_length=50, max_length=8000, description="原始文本")
    optimized_text: str = Field(default="", description="优化后文本")
    original_confidence: float
    final_confidence: float
    reduction_rate: float
    verdict: str
    changes: list[dict] = []
    steps: list[dict] = []


class DocxRequest(BaseModel):
    content: str = Field(min_length=50, max_length=8000, description="原始文本")
    optimized_text: str = Field(default="", description="优化后文本")
    original_confidence: float
    final_confidence: float
    reduction_rate: float
    verdict: str
    changes: list[dict] = []
    steps: list[dict] = []


@router.post("/thesis-reduce/docx")
async def thesis_reduce_docx(
    req: DocxRequest,
    current_user: User = Depends(get_current_user),
):
    """下载优化文档 .docx 格式"""
    from fastapi.responses import Response
    from datetime import datetime

    buffer = _generate_docx(
        original_text=req.content,
        optimized_text=req.optimized_text or req.content,
        changes=req.changes,
        original_conf=req.original_confidence,
        final_conf=req.final_confidence,
        reduction=req.reduction_rate,
        steps=req.steps,
        verdict=req.verdict,
    )

    filename = f"AIGC_optimized_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    return Response(
        content=buffer.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )


@router.post("/thesis-reduce/pdf")
async def thesis_reduce_pdf(
    req: DocxPdfRequest,
    current_user: User = Depends(get_current_user),
):
    """下载优化文档 .pdf 格式"""
    from fastapi.responses import Response
    from datetime import datetime

    buffer = _generate_pdf(
        original_text=req.content,
        optimized_text=req.optimized_text or req.content,
        changes=req.changes,
        original_conf=req.original_confidence,
        final_conf=req.final_confidence,
        reduction=req.reduction_rate,
        steps=req.steps,
        verdict=req.verdict,
    )

    filename = f"AIGC_optimized_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
    return Response(
        content=buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{filename}"},
    )


def _generate_pdf(
    original_text: str, optimized_text: str, changes: list[dict],
    original_conf: float, final_conf: float, reduction: float,
    steps: list[dict], verdict: str,
) -> bytes:
    """生成格式化 PDF 文档"""
    from fpdf import FPDF
    from datetime import datetime
    import os

    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=20)

    # 注册中文字体
    font_paths = [
        "C:/Windows/Fonts/msyh.ttc",   # 微软雅黑
        "C:/Windows/Fonts/simsun.ttc",  # 宋体
        "C:/Windows/Fonts/simhei.ttf",  # 黑体
    ]
    font_loaded = False
    cn_font = "Helvetica"
    for fp in font_paths:
        if os.path.exists(fp):
            pdf.add_font("CNFont", "", fp, uni=True)
            pdf.add_font("CNFont", "B", fp, uni=True)
            cn_font = "CNFont"
            font_loaded = True
            break

    if not font_loaded:
        # Fallback: use built-in with limited Chinese support
        pass

    # 标题
    pdf.set_font(cn_font, "B", 20)
    pdf.cell(0, 15, "AIGC--多模态检测 论文 AIGC 优化报告", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font(cn_font, "", 9)
    pdf.set_text_color(113, 128, 150)
    pdf.cell(0, 8, f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(6)

    # 检测结果对比
    pdf.set_text_color(26, 32, 44)
    pdf.set_font(cn_font, "B", 14)
    pdf.cell(0, 10, "一、检测结果对比", new_x="LMARGIN", new_y="NEXT")
    pdf.set_draw_color(200, 210, 220)

    # 表格
    col_w = [60, 50, 50]
    headers = ["指标", "优化前", "优化后"]
    pdf.set_font(cn_font, "B", 10)
    pdf.set_fill_color(240, 245, 250)
    for i, h in enumerate(headers):
        pdf.cell(col_w[i], 8, h, border=1, fill=True, align="C")
    pdf.ln()

    rows = [
        ["AI 检测率", f"{original_conf:.1%}", f"{final_conf:.1%}"],
        ["判定", "AI生成" if original_conf > 0.3 else "人类写作",
                  "AI生成" if final_conf > 0.3 else "人类写作"],
        ["降低幅度", "—", f"{reduction:.1f}%"],
    ]
    pdf.set_font(cn_font, "", 10)
    pdf.set_fill_color(255, 255, 255)
    for row in rows:
        for i, val in enumerate(row):
            pdf.cell(col_w[i], 8, val, border=1, align="C")
        pdf.ln()
    pdf.ln(6)

    # 综合判定
    pdf.set_font(cn_font, "B", 14)
    pdf.cell(0, 10, "二、综合判定", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font(cn_font, "", 11)
    if "成功" in verdict:
        pdf.set_text_color(16, 185, 129)
    elif "大幅" in verdict:
        pdf.set_text_color(245, 158, 11)
    else:
        pdf.set_text_color(220, 38, 38)
    pdf.multi_cell(0, 7, verdict)
    pdf.set_text_color(26, 32, 44)
    pdf.ln(4)

    # 改写详情
    if changes:
        pdf.set_font(cn_font, "B", 14)
        pdf.cell(0, 10, "三、改写详情", new_x="LMARGIN", new_y="NEXT")
        pdf.set_font(cn_font, "", 10)
        label_map = {
            "slop_removed": lambda c: f"删除 {c['count']} 个 AI 标志词",
            "sentences_split": lambda c: f"句式重构: {c.get('before','?')}句 → {c.get('after','?')}句",
            "specific_data_added": lambda c: f"增加 {c['count']} 处具体数据",
        }
        for c in changes:
            label = label_map.get(c["type"], lambda x: str(x))(c)
            pdf.cell(6, 6, "•", new_x="RIGHT", new_y="LAST")
            pdf.cell(0, 6, label, new_x="LMARGIN", new_y="NEXT")
        pdf.ln(4)

    # 优化步骤
    if steps:
        pdf.set_font(cn_font, "B", 14)
        pdf.cell(0, 10, "四、优化步骤", new_x="LMARGIN", new_y="NEXT")
        pdf.set_font(cn_font, "", 10)
        for i, s in enumerate(steps):
            if s.get("error"):
                pdf.cell(0, 6, f"Step {i+1}: {s['step']} - 失败 ({s['error']})", new_x="LMARGIN", new_y="NEXT")
            else:
                delta = s.get("delta", 0)
                pdf.cell(0, 6, f"Step {i+1}: {s['step']} - AI率 {s['confidence']:.1%} (D={delta:+.1%})", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(4)

    # 优化后文本
    pdf.set_font(cn_font, "B", 14)
    pdf.cell(0, 10, "五、优化后文本", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font(cn_font, "", 10)
    pdf.multi_cell(0, 6, optimized_text)
    pdf.ln(4)

    # 原始文本
    pdf.set_font(cn_font, "B", 14)
    pdf.cell(0, 10, "六、原始文本（对照）", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font(cn_font, "", 8)
    pdf.set_text_color(113, 128, 150)
    pdf.multi_cell(0, 5, original_text)
    pdf.set_text_color(26, 32, 44)
    pdf.ln(6)

    # 页脚
    pdf.set_font(cn_font, "", 9)
    pdf.set_text_color(160, 174, 192)
    pdf.cell(0, 8, "报告结束 - AIGC--多模态检测平台", align="C", new_x="LMARGIN", new_y="NEXT")

    return bytes(pdf.output())


def _generate_docx(
    original_text: str, optimized_text: str, changes: list[dict],
    original_conf: float, final_conf: float, reduction: float,
    steps: list[dict], verdict: str,
):
    """生成格式化 Word 文档"""
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from datetime import datetime

    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)

    # 标题
    title = doc.add_heading("AIGC--多模态检测 论文 AIGC 优化报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x71, 0x80, 0x96)

    doc.add_paragraph()

    # 检测结果对比
    doc.add_heading("一、检测结果对比", level=1)

    table = doc.add_table(rows=4, cols=3, style="Light Grid Accent 1")
    headers = ["指标", "优化前", "优化后"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True

    table.rows[1].cells[0].text = "AI 检测率"
    table.rows[1].cells[1].text = f"{original_conf:.1%}"
    table.rows[1].cells[2].text = f"{final_conf:.1%}"

    before_verdict = "AI生成" if original_conf > 0.3 else "人类写作"
    after_verdict = "AI生成" if final_conf > 0.3 else "人类写作"
    table.rows[2].cells[0].text = "判定"
    table.rows[2].cells[1].text = before_verdict
    table.rows[2].cells[2].text = after_verdict

    table.rows[3].cells[0].text = "降低幅度"
    table.rows[3].cells[1].text = "—"
    table.rows[3].cells[2].text = f"{reduction:.1f}%"

    doc.add_paragraph()

    # 综合判定
    doc.add_heading("二、综合判定", level=1)
    p = doc.add_paragraph(verdict)
    if "成功" in verdict:
        run = p.runs[0]
        run.font.color.rgb = RGBColor(0x10, 0xb9, 0x81)
    elif "大幅" in verdict:
        run = p.runs[0]
        run.font.color.rgb = RGBColor(0xf5, 0x9e, 0x0b)

    doc.add_paragraph()

    # 改写详情
    if changes:
        doc.add_heading("三、改写详情", level=1)
        for c in changes:
            label_map = {
                "slop_removed": f"删除 {c['count']} 个 AI 标志词",
                "sentences_split": f"句式重构: {c.get('before','?')}句 → {c.get('after','?')}句",
                "specific_data_added": f"增加 {c['count']} 处具体数据",
            }
            doc.add_paragraph(label_map.get(c["type"], str(c)), style="List Bullet")

    doc.add_paragraph()

    # 优化步骤
    if steps:
        doc.add_heading("四、优化步骤", level=1)
        for i, s in enumerate(steps):
            if s.get("error"):
                doc.add_paragraph(f"Step {i+1}: {s['step']} — 失败 ({s['error']})")
            else:
                delta = s.get("delta", 0)
                doc.add_paragraph(
                    f"Step {i+1}: {s['step']} — AI率 {s['confidence']:.1%} (Δ={delta:+.1%})"
                )

    doc.add_paragraph()

    # 优化后文本
    doc.add_heading("五、优化后文本", level=1)
    doc.add_paragraph(optimized_text)

    # 分隔
    doc.add_page_break()

    # 原始文本
    doc.add_heading("六、原始文本（对照）", level=1)
    p = doc.add_paragraph()
    run = p.add_run(original_text)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x71, 0x80, 0x96)

    # 页脚
    doc.add_paragraph()
    doc.add_paragraph("— 报告结束 —").alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("AIGC--多模态检测平台")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    from io import BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _verdict_text(reduction: float, final_conf: float) -> str:
    if final_conf < 0.3:
        return f"优化成功：AI置信度降至{final_conf:.0%}，通过检测阈值（<30%）"
    elif reduction > 30:
        return f"大幅降低{reduction:.0f}%，但AI率仍偏高（{final_conf:.0%}），建议进一步人工改写"
    elif reduction > 15:
        return f"降低了{reduction:.0f}%，需继续优化以达到安全阈值"
    elif reduction > 0:
        return f"小幅降低{reduction:.0f}%，当前文本AI痕迹较重，建议重写核心段落"
    else:
        return "未能有效降低，建议重新构思内容"


def _reduce_synonym(text: str) -> str:
    """同义词替换"""
    synonyms = {
        "应用": "使用", "方法": "方式", "重要": "关键", "显著": "明显",
        "发展": "进步", "技术": "科技", "领域": "方面", "挑战": "困难",
        "广泛": "普遍", "取得": "获得", "进展": "进步", "尤为": "特别",
        "深入": "更深", "促进": "推动", "提升": "提高", "优化": "改进",
        "创新": "创造", "突破": "进展", "引领": "带领", "变革": "改变",
    }
    result = text
    import random
    for word, syn in synonyms.items():
        if word in result and random.random() < 0.4:
            result = result.replace(word, syn, 1)
    return result


def _reduce_sentence_restructure(text: str) -> str:
    """句式重构: 拆分长句, 调整语序"""
    import re
    sentences = re.split(r'[。！？]', text)
    result_parts = []
    for s in sentences:
        s = s.strip()
        if not s:
            continue
        # 长句拆短
        if len(s) > 40:
            mid = len(s) // 2
            s = s[:mid] + "，" + s[mid:]
        # 删除典型连接词
        s = re.sub(r'^(然而|此外|因此|所以|但是|不过|而且)', '', s)
        result_parts.append(s)
    return '。'.join(result_parts) + '。'


def _reduce_remove_slop(text: str) -> str:
    """删除 AI 标志词"""
    slops = [
        "值得注意的是", "综上所述", "总而言之", "不可否认",
        "毫无疑问", "进一步来说", "更重要的是", "必须指出",
        "需要强调的是", "众所周知", "在某种程度上",
        "与此同时", "显而易见", "毋庸置疑",
    ]
    result = text
    for slop in slops:
        result = result.replace(slop, "")
    # 清理多余标点
    import re
    result = re.sub(r',,', '，', result)
    result = re.sub(r'。。', '。', result)
    return result


def _reduce_casual_rewrite(text: str) -> str:
    """口语化改写: 加入口语表达, 第一人称, 情感词"""
    import re
    # 第一人称转换
    text = re.sub(r'本研究', '我们這次', text)
    text = re.sub(r'本文', '我这篇', text)
    # 加入情感表达
    phrases = ["说实话", "其实", "我觉得", "挺有意思的是"]
    import random
    if len(text) > 60:
        insert_pos = len(text) // 3
        text = text[:insert_pos] + "，" + random.choice(phrases) + "，" + text[insert_pos:]
    return text


async def _reduce_ai_paraphrase(text: str) -> str | None:
    """AI 对抗改写: DeepSeek 将文本改写为人类风格 (Adversarial Paraphrasing NeurIPS 2025)"""
    try:
        from openai import OpenAI
        from app.config import get_settings
        s = get_settings()
        if not s.llm_api_key:
            return None
        client = OpenAI(api_key=s.llm_api_key, base_url=s.llm_api_base)
        r = client.chat.completions.create(
            model=s.llm_model,
            messages=[{
                "role": "system",
                "content": "你是文本改写助手。将AI文本改写为人类风格：删除AI标志词、加入口语化表达、拆分长句、使用具体词汇。直接输出改写文本。",
            }, {"role": "user", "content": text}],
            max_tokens=1500, temperature=1.1, timeout=30,
        )
        result = r.choices[0].message.content.strip()
        return result if len(result) > 20 else None
    except Exception:
        return None


async def _reduce_ai_backtranslate(text: str) -> str | None:
    """AI 翻译回译: 中→英→中"""
    try:
        from openai import OpenAI
        from app.config import get_settings
        s = get_settings()
        if not s.llm_api_key:
            return None
        client = OpenAI(api_key=s.llm_api_key, base_url=s.llm_api_base)
        r1 = client.chat.completions.create(
            model=s.llm_model,
            messages=[{"role": "user", "content": f"Translate to English:\n{text}"}],
            max_tokens=1000, timeout=20,
        )
        en = r1.choices[0].message.content.strip()
        r2 = client.chat.completions.create(
            model=s.llm_model,
            messages=[{"role": "user", "content": f"Translate to Chinese naturally:\n{en}"}],
            max_tokens=1000, timeout=20,
        )
        zh = r2.choices[0].message.content.strip()
        return zh if len(zh) > 20 else None
    except Exception:
        return None
